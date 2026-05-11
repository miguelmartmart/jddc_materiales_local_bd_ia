"""
server.py — Backend bridge para AI Code Lab (desktop-codelab)

MODO: AI_LOCAL_ONLY — Solo Qwen3 VL 30B en red LAN JDDC.
      NUNCA se conecta a internet. Sin fallback a modelos externos.

Modelos LAN disponibles:
  - jddcia-qwen3-30b-ip  → http://10.13.79.31/api/vlm/v1   (score 100, preferido)
  - jddcia-qwen3-30b     → http://jddcia.local/api/vlm/v1   (score 38, fallback mDNS)

Puerto: 8002 (separado del backend principal en 8001)
"""

import sys
import os
import re
from pathlib import Path
import socket
import asyncio
import json
from typing import Tuple
import httpx

# Add project root to sys.path to allow imports from backend.core
project_root = Path(__file__).parent.parent.parent
sys.path.append(str(project_root))

# Ensure .env exists for pydantic settings
import shutil
env_path = Path(".env")
parent_env_path = project_root / ".env"

if not env_path.exists():
    if parent_env_path.exists():
        print(f"Copying .env from {parent_env_path} to {env_path}")
        shutil.copy(parent_env_path, env_path)
    else:
        print(f"Warning: .env not found in {env_path} or {parent_env_path}")

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List, Dict, Any
import uvicorn
import logging

# Import existing AI components
try:
    from backend.core.factory.ai_factory import AIFactory
    from backend.core.abstract.ai import AIConfig
except ImportError as e:
    print(f"Error importing backend modules: {e}")
    print(f"Sys path: {sys.path}")
    sys.exit(1)

# Modular prompt engine — loaded after all setup so errors are surfaced clearly
_backend_dir = Path(__file__).parent
sys.path.insert(0, str(_backend_dir))
try:
    from prompt_engine import PromptBuilder, PhaseContext, DEFAULT_PHASES
    _PROMPT_BUILDER = PromptBuilder(DEFAULT_PHASES)
    print(f"[prompt_engine] Loaded {len(DEFAULT_PHASES)} phases: "
          f"{', '.join(p.name for p in DEFAULT_PHASES)}")
except Exception as _pe_err:
    print(f"[prompt_engine] WARNING: Failed to load — falling back to legacy prompt. "
          f"Error: {_pe_err}")
    _PROMPT_BUILDER = None  # type: ignore
    PhaseContext = None  # type: ignore

# ─── MODELOS LAN HARDCODEADOS (temporal hasta que se implemente ModelManager) ───
class LocalModelIds:
    ALL = frozenset([
        "jddcia-qwen3-30b-ip",    # IP directa (preferido)
        "jddcia-qwen3-30b"        # mDNS fallback
    ])

class ModelManager:
    def list_models(self, enabled_only=True):
        """Devuelve modelos LAN hardcodeados con las nuevas URLs."""
        import os

        # Leer variables de entorno para URLs base
        base_url_ip = os.environ.get('JDDCIA_BASE_URL', 'http://10.13.79.31/api/vlm/v1')
        base_url_mdns = os.environ.get('JDDCIA_BASE_URL_FALLBACK', 'http://jddcia.local/api/vlm/v1')

        models = [
            {
                "id": "jddcia-qwen3-30b-ip",
                "name": "Qwen3 VL 30B (IP directa)",
                "model_id": "unified-main",
                "provider": "openai",
                "schema": "openai",
                "base_url": base_url_ip,
                "api_key": "jddcia-local-key",
                "headers": {
                    "Authorization": "Basic YWRtaW46YWlzdGFjazIwMjY="
                },
                "enabled": True,
                "score": 100,
                "description": "Modelo Qwen3 VL 30B local vía IP directa"
            },
            {
                "id": "jddcia-qwen3-30b",
                "name": "Qwen3 VL 30B (mDNS)",
                "model_id": "unified-main",
                "provider": "openai",
                "schema": "openai",
                "base_url": base_url_mdns,
                "api_key": "jddcia-local-key",
                "headers": {
                    "Authorization": "Basic YWRtaW46YWlzdGFjazIwMjY="
                },
                "enabled": True,
                "score": 38,
                "description": "Modelo Qwen3 VL 30B local vía mDNS"
            }
        ]

        if enabled_only:
            models = [m for m in models if m.get('enabled', True)]

        return models

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger("CodeLabBackend")

# ─── Constantes LAN ───────────────────────────────────────────────────────────
# IDs de modelos locales (fuente única de verdad: LocalModelIds)
LOCAL_MODEL_IDS: frozenset = LocalModelIds.ALL

# Timeout para Qwen3 30B (puede tardar 30-60s en respuestas largas de código)
LAN_READ_TIMEOUT_S = 180

# ─── AUTODESCUBRIMIENTO Y TAILSCALE ───────────────────────────────────────────
# Rutas de conexión al servidor JDDC (en orden de prioridad)
JDDC_DISCOVERY_ROUTES = [
    # Tailscale VPN (más confiable en DHCP)
    {"name": "Tailscale VPN", "host": "100.73.44.46", "ports": [8002, 8001, 80], "priority": 1},
    # Nueva IP LAN del servidor IA (después de configuración firewall)
    {"name": "LAN IP (10.13.79.31)", "host": "10.13.79.31", "ports": [80, 8001, 8002], "priority": 2},
    # IP local anterior (por compatibilidad)
    {"name": "IP Local (192.168.0.36)", "host": "192.168.0.36", "ports": [8002, 8001, 80], "priority": 3},
    # DNS local (mDNS) - ahora con puertos de IA prioritarios
    {"name": "mDNS Local", "host": "jddcia.local", "ports": [80, 8001, 8002], "priority": 4},
]

# Caché de última IP exitosa (se carga/guarda en .jddcia_cache.json)
_CACHE_FILE = Path(".jddcia_cache.json")
_last_working_route: dict = {}

def _load_cache():
    """Carga la última ruta de conexión exitosa desde caché."""
    global _last_working_route
    try:
        if _CACHE_FILE.exists():
            with open(_CACHE_FILE, "r") as f:
                _last_working_route = json.load(f)
                logger.info(f"[AutoDiscovery] Caché cargado: {_last_working_route.get('name', 'desconocido')}")
    except Exception as e:
        logger.warning(f"[AutoDiscovery] No se pudo cargar caché: {e}")
        _last_working_route = {}

def _save_cache(route: dict):
    """Guarda la ruta de conexión exitosa en caché."""
    global _last_working_route
    try:
        _last_working_route = route
        with open(_CACHE_FILE, "w") as f:
            json.dump(route, f, indent=2)
        logger.info(f"[AutoDiscovery] Caché guardado: {route.get('name', 'desconocido')}")
    except Exception as e:
        logger.warning(f"[AutoDiscovery] No se pudo guardar caché: {e}")

async def _probe_route(route: dict, timeout_s: float = 1.0) -> Tuple[bool, str]:
    """
    Verifica si una ruta de conexión es accesible.
    Busca cualquier endpoint JDDC disponible, priorizando health checks.
    Devuelve la URL del puerto donde encontró el servidor.
    """
    host = route.get("host")
    ports = route.get("ports", [8002])
    name = route.get("name", host)
    
    # Primero buscar health checks (más rápidos y confiables)
    for port in ports:
        url = f"http://{host}:{port}/health"
        try:
            async with httpx.AsyncClient(timeout=timeout_s) as client:
                response = await client.get(url)
                if response.status_code == 200:
                    logger.info(f"[AutoDiscovery] ✅ {name} accesible en {url}")
                    # Usar el mismo puerto donde encontró el servidor
                    return True, f"http://{host}:{port}"
        except Exception as e:
            logger.debug(f"[AutoDiscovery] {name}:{port}/health → {type(e).__name__}")
    
    # Si no hay health checks, probar endpoints de IA directamente
    for port in [8001, 8002, 80]:
        url = f"http://{host}:{port}/api/vlm/v1/models"
        try:
            async with httpx.AsyncClient(timeout=timeout_s) as client:
                response = await client.get(url)
                if response.status_code == 200:
                    logger.info(f"[AutoDiscovery] ✅ {name} IA accesible en {url}")
                    return True, f"http://{host}:{port}"
        except Exception as e:
            logger.debug(f"[AutoDiscovery] {name}:{port}/api/vlm/v1/models → {type(e).__name__}")
    
    return False, ""

async def _autodiscover_server() -> dict:
    """
    Intenta encontrar la mejor ruta de conexión al servidor JDDC.
    Estrategia: Buscar rápidamente cualquier servidor JDDC disponible.
    """
    global _last_working_route
    
    logger.info("[AutoDiscovery] Iniciando autodescubrimiento rápido de servidor JDDC...")
    
    # 1. Intentar caché primero (timeout muy corto)
    if _last_working_route:
        success, url = await _probe_route(_last_working_route, timeout_s=0.5)
        if success:
            logger.info(f"[AutoDiscovery] ✅ Ruta en caché funciona: {url}")
            return {"host": _last_working_route["host"], "url": url, "source": "cache"}
    
    # 2. Intentar rutas en orden de prioridad (timeout agresivo)
    sorted_routes = sorted(JDDC_DISCOVERY_ROUTES, key=lambda r: r.get("priority", 999))
    for route in sorted_routes:
        logger.info(f"[AutoDiscovery] Probando ruta: {route.get('name', route.get('host'))}")
        success, url = await _probe_route(route, timeout_s=0.5)  # Timeout muy corto
        if success:
            # Si encontramos una ruta, guardamos la ruta completa
            _save_cache(route)
            return {"host": route["host"], "url": url, "source": route["name"]}

    # 3. Si nada funciona, devolver configuración por defecto (IP local)
    logger.warning("[AutoDiscovery] No se encontró servidor accesible. Usando configuración por defecto.")
    default_route = JDDC_DISCOVERY_ROUTES[1]  # IP local por defecto
    return {
        "host": default_route["host"],
        "url": f"http://{default_route['host']}:80",  # Puerto 80 por defecto
        "source": "default"
    }


class GenerateRequest(BaseModel):
    prompt: str
    messages: Optional[List[Dict[str, str]]] = None
    code_context: Optional[str] = None
    model_id: Optional[str] = None
    project_path: Optional[str] = None
    system_prompt: Optional[str] = None
    images: Optional[List[str]] = None  # Lista de imágenes en base64 (data:image/...;base64,...)
    documents: Optional[List[Dict[str, str]]] = None  # [{"name": "file.pdf", "text": "..."}]


class ExtractTextRequest(BaseModel):
    filename: str          # Nombre original del archivo (para detectar tipo)
    content_b64: str       # Contenido del archivo en base64


class ExtractTextResponse(BaseModel):
    filename: str
    text: str              # Texto extraído
    pages: Optional[int] = None
    rows: Optional[int] = None
    truncated: bool = False
    error: Optional[str] = None


class ModelListResponse(BaseModel):
    models: List[Dict[str, Any]]


# ─── Cargar caché al iniciar
_load_cache()

app = FastAPI(
    title="AI Code Lab Backend",
    description="Backend bridge para CodeLab — Solo Qwen3 LAN, sin internet (con Tailscale autodetección)."
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class GenerateRequest(BaseModel):
    prompt: str
    messages: Optional[List[Dict[str, str]]] = None
    code_context: Optional[str] = None
    model_id: Optional[str] = None
    project_path: Optional[str] = None
    system_prompt: Optional[str] = None
    images: Optional[List[str]] = None  # Lista de imágenes en base64 (data:image/...;base64,...)
    documents: Optional[List[Dict[str, str]]] = None  # [{"name": "file.pdf", "text": "..."}]


class ExtractTextRequest(BaseModel):
    filename: str          # Nombre original del archivo (para detectar tipo)
    content_b64: str       # Contenido del archivo en base64


class ExtractTextResponse(BaseModel):
    filename: str
    text: str              # Texto extraído
    pages: Optional[int] = None
    rows: Optional[int] = None
    truncated: bool = False
    error: Optional[str] = None


class ModelListResponse(BaseModel):
    models: List[Dict[str, Any]]


def _get_lan_models() -> List[Dict[str, Any]]:
    """
    Devuelve SOLO los modelos LAN habilitados, ordenados por score descendente.
    Filtra cualquier modelo que no sea local (sin internet).
    """
    try:
        manager = ModelManager()
        all_models = manager.list_models(enabled_only=True) or []
        lan_models = [m for m in all_models if m.get('id') in LOCAL_MODEL_IDS]
        # Ordenar por score descendente (mayor score = más prioritario)
        lan_models.sort(key=lambda m: m.get('score', 0), reverse=True)
        return lan_models
    except Exception as e:
        logger.error(f"Error obteniendo modelos LAN: {e}")
        return []


@app.get("/health")
async def health():
    """Health check — verifica que el backend está corriendo."""
    lan_models = _get_lan_models()
    return {
        "status": "ok",
        "mode": "AI_LOCAL_ONLY",
        "lan_models": [m.get('name') for m in lan_models],
        "lan_models_count": len(lan_models)
    }


@app.get("/debug/ping-ai")
async def debug_ping_ai():
    """
    Diagnóstico: envía un mensaje mínimo a cada modelo LAN y devuelve
    el resultado completo (status HTTP, body, excepción si la hay).
    Abre http://localhost:8002/debug/ping-ai en el navegador para ver el error exacto.
    """
    import traceback as _tb
    results = []
    for model in _get_lan_models():
        url = f"{model['base_url']}/chat/completions"
        headers = model.get('headers', {})
        payload = {
            "model": model.get('model_id', 'unified-main'),
            "messages": [{"role": "user", "content": "hola"}],
            "max_tokens": 10,
        }
        entry: Dict[str, Any] = {"model": model['id'], "url": url}
        try:
            async with httpx.AsyncClient(timeout=15.0, trust_env=False) as client:
                resp = await client.post(url, json=payload, headers=headers)
                entry["status"] = resp.status_code
                entry["body"] = resp.text[:2000]
        except Exception as exc:
            entry["error"] = f"{type(exc).__name__}: {exc}"
            entry["traceback"] = _tb.format_exc()
        results.append(entry)
    return {"results": results}


@app.get("/api/models", response_model=ModelListResponse)
async def list_models():
    """
    Lista SOLO los modelos LAN disponibles.
    No incluye modelos de internet (Groq, Gemini, OpenAI, etc.).
    """
    try:
        lan_models = _get_lan_models()
        if not lan_models:
            logger.warning("No se encontraron modelos LAN habilitados. "
                           "Verifica que jddcia-qwen3-30b esté enabled=true en jddcia_models.json")
        logger.info(f"Modelos LAN disponibles: {[m.get('name') for m in lan_models]}")
        return {"models": lan_models}
    except Exception as e:
        logger.error(f"Error listando modelos: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/models/refresh")
async def refresh_models():
    """
    En modo AI_LOCAL_ONLY, no se conecta a APIs externas.
    Solo verifica el estado de los modelos LAN.
    """
    lan_models = _get_lan_models()
    return {
        "message": f"Modo AI_LOCAL_ONLY activo. {len(lan_models)} modelos LAN disponibles: "
                   f"{[m.get('name') for m in lan_models]}. "
                   f"No se conecta a internet."
    }


@app.get("/api/autodiscover")
async def autodiscover():
    """
    Endpoint de autodescubrimiento: devuelve configuración por defecto inmediatamente
    y hace autodescubrimiento en background para futuras peticiones.
    
    Respuesta inmediata:
    {
      "status": "ok",
      "host": "jddcia.local",
      "url": "http://jddcia.local:80",
      "source": "default",
      "lan_models": [...]
    }
    """
    lan_models = _get_lan_models()
    
    # Devolver configuración por defecto inmediatamente
    default_config = {
        "status": "ok",
        "host": "jddcia.local",
        "url": "http://jddcia.local:80",
        "source": "default",
        "lan_models": [m.get('name') for m in lan_models],
        "lan_models_count": len(lan_models),
        "routes_available": len(JDDC_DISCOVERY_ROUTES)
    }
    
    # Hacer autodescubrimiento en background (no bloquea la respuesta)
    asyncio.create_task(_background_autodiscover())
    
    return default_config


async def _background_autodiscover():
    """Hace autodescubrimiento en background sin bloquear respuestas."""
    try:
        await _autodiscover_server()
        logger.info("[AutoDiscovery] Autodescubrimiento en background completado")
    except Exception as e:
        logger.warning(f"[AutoDiscovery] Error en autodescubrimiento background: {e}")


@app.post("/api/autodiscover/verify")
async def autodiscover_verify(request: dict):
    """
    Verifica si una ruta específica es accesible.
    Útil para validar antes de usar una conexión.
    
    Request: {"host": "100.73.44.46", "port": 8002}
    """
    try:
        host = request.get("host", "127.0.0.1")
        port = request.get("port", 8002)
        timeout = request.get("timeout_s", 2.0)
        
        url = f"http://{host}:{port}/health"
        async with httpx.AsyncClient(timeout=timeout) as client:
            response = await client.get(url)
            if response.status_code == 200:
                logger.info(f"[Verify] ✅ {host}:{port} accesible")
                return {
                    "status": "ok",
                    "host": host,
                    "port": port,
                    "url": url,
                    "accessible": True
                }
            else:
                logger.warning(f"[Verify] ⚠️ {host}:{port} devuelve {response.status_code}")
                return {
                    "status": "error",
                    "host": host,
                    "port": port,
                    "accessible": False,
                    "error": f"HTTP {response.status_code}"
                }
    except Exception as e:
        logger.error(f"[Verify] ❌ Error verificando {host}:{port}: {e}")
        return {
            "status": "error",
            "host": host,
            "port": port,
            "accessible": False,
            "error": str(e)
        }


async def _try_lan_model(
    model_config: Dict[str, Any],
    prompt: str,
    system_prompt: str,
    attempt: int = 1,
    images: Optional[List[str]] = None
) -> Optional[str]:
    """
    Intenta generar respuesta con un modelo LAN específico usando HTTP directo.
    Devuelve la respuesta o None si falla.
    
    Usa httpx para autenticación Basic Auth correcta.
    """
    import base64
    import httpx
    
    model_name = model_config.get('name', 'Unknown')
    model_id = model_config.get('id', '')
    base_url = model_config.get('base_url', '')

    # Verificar que es un modelo LAN (seguridad adicional)
    if model_id not in LOCAL_MODEL_IDS:
        logger.error(f"[LAN_ONLY] Modelo {model_id} NO es local. Bloqueado.")
        return None

    try:
        has_images = bool(images)
        logger.info(f"[LAN_ONLY] Intentando {model_name} (intento {attempt})"
                    + (f" con {len(images)} imagen(es)" if has_images else ""))

        if not base_url:
            logger.error(f"[LAN_ONLY] Sin base_url para {model_name}")
            return None

        # Preparar headers con autenticación Basic Auth
        headers = {}
        auth_header = model_config.get('headers', {}).get('Authorization')
        if auth_header:
            headers['Authorization'] = auth_header
            logger.info(f"[LAN_ONLY] Usando autenticación Basic Auth")
        
        # Preparar payload para chat/completions
        messages = []
        if system_prompt:
            messages.append({"role": "system", "content": system_prompt})
        
        # Agregar imágenes si las hay (formato OpenAI Vision)
        if has_images and images:
            for img in images:
                if not img.startswith("data:"):
                    img = f"data:image/jpeg;base64,{img}"
                messages.append({"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": img}}
                ]})
        
        messages.append({"role": "user", "content": prompt})
        
        # Contexto total del modelo: 4096 tokens (input + output).
        # Probamos con 2000 tokens de salida; si la API devuelve 400 por overflow
        # reducimos hasta que encaje.
        max_tokens_candidates = [2000, 1500, 1000, 512]

        # trust_env=False: ignora HTTP_PROXY/HTTPS_PROXY del entorno de Windows/Electron.
        # Las IPs LAN (10.x.x.x) deben ir directas, no por ningún proxy.
        async with httpx.AsyncClient(timeout=LAN_READ_TIMEOUT_S, trust_env=False) as client:
            chat_url = f"{base_url}/chat/completions"
            logger.info(f"[LAN_ONLY] POST {chat_url} | sys_prompt={len(system_prompt)}chars")

            last_error = ""
            for max_tok in max_tokens_candidates:
                payload = {
                    "model": model_config.get('model_id', 'unified-main'),
                    "messages": messages,
                    "temperature": 0.7,
                    "max_tokens": max_tok
                }
                resp = await client.post(chat_url, json=payload, headers=headers)

                if resp.status_code == 400:
                    last_error = resp.text[:500]
                    logger.warning(f"[LAN_ONLY] HTTP 400 (max_tokens={max_tok}): {last_error}")
                    _overflow_kws = (
                        "max_tokens", "context length", "context_length",
                        "prompt", "token", "length", "exceed", "too long",
                        "input", "maximum",
                    )
                    if any(kw in last_error.lower() for kw in _overflow_kws):
                        logger.warning(f"[LAN_ONLY] Parece overflow — reintentando con menos tokens.")
                        continue
                    # 400 no recuperable (auth, formato, etc.)
                    logger.error(f"[LAN_ONLY] Error 400 no recuperable: {last_error}")
                    return None

                if resp.status_code != 200:
                    logger.error(f"[LAN_ONLY] Error HTTP {resp.status_code}: {resp.text[:500]}")
                    return None

                data = resp.json()
                response_text = data.get('choices', [{}])[0].get('message', {}).get('content', '')

                if response_text:
                    logger.info(f"[LAN_ONLY] Respuesta de {model_name} (max_tokens={max_tok})"
                                + (" (vision)" if has_images else ""))
                    return response_text
                else:
                    logger.warning(f"[LAN_ONLY] Respuesta vacia de {model_name}")
                    return None

            # Último recurso: calcular los tokens exactos disponibles desde el mensaje de error.
            # El error dice: "request has N input tokens" y "context length is M".
            # Si quedan al menos 30 tokens libres, intentamos con ese valor exacto.
            _m_input = re.search(r'has (\d+) input tokens', last_error)
            _m_ctx   = re.search(r'context length is (\d+)', last_error)
            if _m_input and _m_ctx:
                _avail = int(_m_ctx.group(1)) - int(_m_input.group(1)) - 5
                if _avail >= 30:
                    logger.warning(f"[LAN_ONLY] Último intento con max_tokens exactos: {_avail}")
                    payload["max_tokens"] = _avail
                    resp_last = await client.post(chat_url, json=payload, headers=headers)
                    if resp_last.status_code == 200:
                        _rt = resp_last.json().get('choices', [{}])[0].get('message', {}).get('content', '')
                        if _rt:
                            logger.info(f"[LAN_ONLY] Respuesta corta con {_avail} tokens: OK")
                            return _rt
                else:
                    logger.error(
                        f"[LAN_ONLY] Input demasiado grande: quedan solo {_avail} tokens libres. "
                        f"Reducir el system prompt."
                    )

            logger.error(f"[LAN_ONLY] Todos los intentos de max_tokens fallaron: {last_error}")
            return None

    except Exception as e:
        logger.error(f"[LAN_ONLY] Error con {model_name}: {type(e).__name__}: {e}")
        import traceback
        logger.error(f"[LAN_ONLY] Traceback: {traceback.format_exc()}")
        return None


# ─── Extracción de texto de documentos ───────────────────────────────────────

MAX_DOC_CHARS = 12000  # Límite de caracteres por documento (para no saturar el contexto)

def _extract_text_from_bytes(filename: str, data: bytes) -> ExtractTextResponse:
    """
    Extrae texto de un documento según su extensión.
    Soporta: PDF, DOCX, XLSX, XLS, CSV, TXT, MD, JSON, XML, HTML, PY, JS, TS, etc.
    """
    ext = Path(filename).suffix.lower()
    text = ""
    pages = None
    rows = None
    error = None

    try:
        # ── PDF ──────────────────────────────────────────────────────────────
        if ext == ".pdf":
            try:
                import PyPDF2
                import io
                reader = PyPDF2.PdfReader(io.BytesIO(data))
                pages = len(reader.pages)
                parts = []
                for i, page in enumerate(reader.pages):
                    try:
                        t = page.extract_text() or ""
                        if t.strip():
                            parts.append(f"[Página {i+1}]\n{t}")
                    except Exception:
                        pass
                text = "\n\n".join(parts)
            except ImportError:
                error = "PyPDF2 no instalado. Instala con: pip install PyPDF2"

        # ── DOCX (Word) ───────────────────────────────────────────────────────
        elif ext == ".docx":
            try:
                import docx
                import io
                doc = docx.Document(io.BytesIO(data))
                parts = [p.text for p in doc.paragraphs if p.text.strip()]
                # También extraer tablas
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            parts.append(row_text)
                text = "\n".join(parts)
            except ImportError:
                error = "python-docx no instalado. Instala con: pip install python-docx"

        # ── XLSX (Excel moderno) ──────────────────────────────────────────────
        elif ext == ".xlsx":
            try:
                import openpyxl
                import io
                wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
                parts = []
                total_rows = 0
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    sheet_rows = []
                    for row in ws.iter_rows(values_only=True):
                        row_vals = [str(c) if c is not None else "" for c in row]
                        if any(v.strip() for v in row_vals):
                            sheet_rows.append(" | ".join(row_vals))
                            total_rows += 1
                    if sheet_rows:
                        parts.append(f"[Hoja: {sheet_name}]\n" + "\n".join(sheet_rows))
                text = "\n\n".join(parts)
                rows = total_rows
            except ImportError:
                error = "openpyxl no instalado. Instala con: pip install openpyxl"

        # ── XLS (Excel antiguo) ───────────────────────────────────────────────
        elif ext == ".xls":
            try:
                import xlrd
                import io
                wb = xlrd.open_workbook(file_contents=data)
                parts = []
                total_rows = 0
                for sheet in wb.sheets():
                    sheet_rows = []
                    for rx in range(sheet.nrows):
                        row_vals = [str(sheet.cell_value(rx, cx)) for cx in range(sheet.ncols)]
                        if any(v.strip() for v in row_vals):
                            sheet_rows.append(" | ".join(row_vals))
                            total_rows += 1
                    if sheet_rows:
                        parts.append(f"[Hoja: {sheet.name}]\n" + "\n".join(sheet_rows))
                text = "\n\n".join(parts)
                rows = total_rows
            except ImportError:
                error = "xlrd no instalado. Instala con: pip install xlrd"

        # ── CSV ───────────────────────────────────────────────────────────────
        elif ext == ".csv":
            import csv
            import io
            # Intentar detectar encoding
            try:
                decoded = data.decode("utf-8")
            except UnicodeDecodeError:
                decoded = data.decode("latin-1", errors="replace")
            reader = csv.reader(io.StringIO(decoded))
            csv_rows = list(reader)
            rows = len(csv_rows)
            text = "\n".join(" | ".join(row) for row in csv_rows if any(c.strip() for c in row))

        # ── Texto plano / código / markdown / JSON / XML / HTML ───────────────
        elif ext in (".txt", ".md", ".json", ".xml", ".html", ".htm",
                     ".py", ".js", ".ts", ".jsx", ".tsx", ".css", ".scss",
                     ".yaml", ".yml", ".toml", ".ini", ".cfg", ".env",
                     ".sh", ".bat", ".ps1", ".sql", ".log"):
            try:
                text = data.decode("utf-8")
            except UnicodeDecodeError:
                text = data.decode("latin-1", errors="replace")

        else:
            # Intentar como texto plano como último recurso
            try:
                text = data.decode("utf-8")
            except Exception:
                error = f"Formato '{ext}' no soportado directamente. Tipos soportados: PDF, DOCX, XLSX, XLS, CSV, TXT, MD, JSON, XML, HTML, y archivos de código."

    except Exception as e:
        error = f"Error extrayendo texto: {type(e).__name__}: {e}"
        logger.error(f"[ExtractText] Error en {filename}: {error}")

    # Truncar si es muy largo
    truncated = False
    if len(text) > MAX_DOC_CHARS:
        text = text[:MAX_DOC_CHARS] + f"\n\n[... DOCUMENTO TRUNCADO — {len(text)} caracteres totales, mostrando primeros {MAX_DOC_CHARS} ...]"
        truncated = True

    return ExtractTextResponse(
        filename=filename,
        text=text,
        pages=pages,
        rows=rows,
        truncated=truncated,
        error=error
    )


@app.post("/api/extract-text", response_model=ExtractTextResponse)
async def extract_text(request: ExtractTextRequest):
    """
    Extrae texto de un documento enviado en base64.
    El frontend envía el archivo como base64 y recibe el texto extraído.
    El texto se usa como contexto adicional en el prompt de la IA.
    
    Formatos soportados:
    - PDF (.pdf) — via PyPDF2
    - Word (.docx) — via python-docx
    - Excel (.xlsx) — via openpyxl
    - Excel antiguo (.xls) — via xlrd
    - CSV (.csv) — nativo Python
    - Texto plano: .txt, .md, .json, .xml, .html, .py, .js, .ts, etc.
    """
    import base64
    try:
        # Decodificar base64 (puede venir con o sin prefijo data:...)
        b64 = request.content_b64
        if "," in b64:
            b64 = b64.split(",", 1)[1]
        data = base64.b64decode(b64)
        
        result = _extract_text_from_bytes(request.filename, data)
        
        if result.error and not result.text:
            raise HTTPException(status_code=422, detail=result.error)
        
        logger.info(
            f"[ExtractText] ✅ {request.filename} → "
            f"{len(result.text)} chars"
            + (f", {result.pages} páginas" if result.pages else "")
            + (f", {result.rows} filas" if result.rows else "")
            + (" (truncado)" if result.truncated else "")
        )
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[ExtractText] Error inesperado: {type(e).__name__}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/generate")
async def generate_code(request: GenerateRequest):
    """
    Genera código usando SOLO Qwen3 LAN.
    
    Estrategia:
    1. Intenta con jddcia-qwen3-30b-ip (IP directa, score 95)
    2. Si falla, intenta con jddcia-qwen3-30b (mDNS, score 50)
    3. Si ambos fallan, devuelve error 503 (sin fallback a internet)
    """
    try:
        logger.info(f"[CodeLab] Nueva petición de generación. Modelo solicitado: {request.model_id or 'Auto LAN'}")

        # Obtener modelos LAN disponibles
        lan_models = _get_lan_models()
        if not lan_models:
            raise HTTPException(
                status_code=503,
                detail="No hay modelos LAN disponibles. Verifica que jddcia-qwen3-30b esté habilitado "
                       "en backend/.env y que el servidor Qwen3 "
                       "esté accesible en 10.13.79.31 o jddcia.local."
            )

        # Si se solicita un modelo específico, verificar que sea LAN
        if request.model_id:
            if request.model_id not in LOCAL_MODEL_IDS:
                logger.warning(
                    f"[LAN_ONLY] Modelo solicitado '{request.model_id}' no es local. "
                    f"Usando el mejor modelo LAN disponible."
                )
            else:
                # Mover el modelo solicitado al principio si es LAN
                requested = next((m for m in lan_models if m.get('id') == request.model_id), None)
                if requested:
                    lan_models = [requested] + [m for m in lan_models if m.get('id') != request.model_id]

        # Construir contexto del proyecto
        base_folder = "."
        existing_folders = []
        if request.project_path:
            base_folder = os.path.basename(os.path.normpath(request.project_path))
            try:
                with os.scandir(request.project_path) as entries:
                    existing_folders = [
                        entry.name for entry in entries
                        if entry.is_dir() and not entry.name.startswith('.')
                    ]
            except Exception as e:
                logger.warning(f"No se pudo listar directorios en {request.project_path}: {e}")

        # Construir system prompt si no se proporcionó
        if not request.system_prompt:
            has_images = bool(request.images)
            has_documents = bool(request.documents)
            folders_context = ", ".join(existing_folders) if existing_folders else "(Ninguna detectada)"

            # Attachment summary for the header phase
            content_types = []
            if has_images:
                content_types.append(f"{len(request.images)} imagen(es)")
            if has_documents:
                content_types.append(f"{len(request.documents)} documento(s)")
            content_desc = (" con " + " y ".join(content_types) + " adjunto(s)") if content_types else ""

            attachments_context = ""
            if has_images:
                n = len(request.images)
                attachments_context += (
                    f"\n**IMÁGENES ADJUNTAS ({n} en total)**: "
                    "Debes analizarlas visualmente. Son adjuntos visuales reales, NO los ignores.\n"
                )
            if has_documents:
                names = [d.get("name", "sin nombre") for d in request.documents]
                attachments_context += (
                    f"\n**DOCUMENTOS ADJUNTOS ({len(names)} en total)**:\n"
                    + "\n".join(f"  {i+1}. {n}" for i, n in enumerate(names))
                    + "\n"
                )

            if _PROMPT_BUILDER is not None:
                ctx = PhaseContext(
                    base_folder=base_folder,
                    project_path=request.project_path,
                    folders_context=folders_context,
                    has_images=has_images,
                    has_documents=has_documents,
                    attachments_context=attachments_context,
                    content_desc=content_desc,
                )
                request.system_prompt = _PROMPT_BUILDER.build(ctx)
            else:
                # Fallback mínimo si prompt_engine no cargó
                request.system_prompt = (
                    f"Eres un asistente de programación trabajando en el proyecto '{base_folder}'. "
                    "Responde siempre en español. Genera código completo y funcional."
                )

            # ── Truncación de seguridad ───────────────────────────────────────
            # El modelo LAN tiene 4096 tokens totales (input + output).
            # Con ~4 chars/token, 8000 chars ≈ 2000 tokens, dejando ~2096 para
            # historial de conversación + respuesta del modelo.
            _SYS_CHAR_LIMIT = 8000
            if len(request.system_prompt) > _SYS_CHAR_LIMIT:
                cut = request.system_prompt.rfind('\n\n', 0, _SYS_CHAR_LIMIT)
                if cut < 500:
                    cut = _SYS_CHAR_LIMIT
                request.system_prompt = request.system_prompt[:cut]
                logger.warning(
                    f"[CodeLab] System prompt truncado: {len(request.system_prompt)} chars "
                    f"(límite {_SYS_CHAR_LIMIT}, ventana modelo 4096 tokens)"
                )

        # Construir prompt completo con historial
        is_analysis_mode = bool(request.images or request.documents)

        def _strip_code_structure(text: str) -> str:
            """Elimina Resumen, Próximos Pasos y [[NEXT_ACTION]] de respuestas conversacionales."""
            # Quitar [[NEXT_ACTION:...]] con re.DOTALL para que . cruce líneas
            text = re.sub(r'\[\[\s*NEXT_ACTION\s*:.*?\]\]', '', text, flags=re.IGNORECASE | re.DOTALL)
            # Cortar el texto en el primer marcador de sección innecesaria
            lower = text.lower()
            cut_at = len(text)
            for marker in ['## resumen', '## próximos pasos', '## proximos pasos']:
                pos = lower.find(marker)
                if pos != -1 and pos < cut_at:
                    cut_at = pos
            text = text[:cut_at]
            return text.rstrip()

        def _is_code_task(prompt: str) -> bool:
            """Devuelve True si el mensaje parece una tarea de código/acción, no conversacional.

            Usa word boundaries para evitar falsos positivos como "capital" → "api".
            """
            if not prompt:
                return False
            p = prompt.lower().strip()
            # Palabras clave que deben aparecer como palabras completas (word boundary)
            code_words = [
                "crea", "cree", "genera", "escribe", "implementa",
                "añade", "agrega", "modifica", "cambia", "actualiza",
                "elimina", "borra", "instala", "ejecuta", "compila",
                "construye", "refactoriza", "arregla", "corrige", "soluciona",
                "depura", "debugea",
                "script", "función", "clase", "archivo", "fichero", "módulo",
                "carpeta", "directorio", "estructura", "proyecto",
                "componente", "api", "endpoint", "database", "sql",
                "python", "javascript", "typescript", "html", "css", "react",
                "import", "function", "mkdir", "powershell", "terminal",
                "programa", "aplicación", "app",
            ]
            pattern = r'\b(?:' + '|'.join(re.escape(w) for w in code_words) + r')\b'
            if re.search(pattern, p, re.IGNORECASE):
                return True
            # Snippets de código que no requieren word boundary
            if any(snippet in p for snippet in ["def ", "class ", "const ", "let ", "var ", "base de datos"]):
                return True
            # Mensajes largos sin ninguna de las anteriores son probablemente tareas
            if len(p) > 120:
                return True
            return False
        # ─── Truncar historial para evitar overflow de 4096 tokens ──────────────
        # El modelo tiene 4096 tokens totales. El system prompt usa ~1600 tokens.
        # Quedan ~2496 para historial + respuesta. Sin truncación, conversaciones
        # largas dejan <500 tokens para la respuesta → IA miente o trunca respuestas.
        #
        # Estrategia:
        # 1. Mensajes [SYSTEM] con salidas de terminal se truncan a 300 chars
        #    (la IA solo necesita saber éxito/error, no el listado completo de dirs)
        # 2. Solo se envían los últimos MAX_HIST_MESSAGES mensajes al modelo
        _MAX_HIST_MSG   = 10    # máximo de mensajes del historial enviados al modelo
        _MAX_SYS_CHARS  = 300   # máximo de chars por mensaje [SYSTEM]

        def _truncate_history(msgs):
            if not msgs:
                return msgs
            recent = msgs[-_MAX_HIST_MSG:]
            result = []
            for m in recent:
                if m.get("role") == "system" and len(m.get("content", "")) > _MAX_SYS_CHARS:
                    short = m["content"][:_MAX_SYS_CHARS].rstrip() + "…[truncado]"
                    result.append({**m, "content": short})
                else:
                    result.append(m)
            return result

        full_prompt = ""
        if request.messages:
            for msg in _truncate_history(request.messages):
                role = msg.get("role", "user").upper()
                content = msg.get("content", "")
                full_prompt += f"{role}: {content}\n\n"
            full_prompt += f"USER: {request.prompt}\n\n"
            if not is_analysis_mode and _is_code_task(request.prompt):
                full_prompt += "\n\n**INSTRUCCIÓN MAESTRA FINAL (CRÍTICA)**:\n"
                full_prompt += "1. Genera TODO el código solicitado primero.\n"
                full_prompt += "2. AL FINAL, añade obligatoriamente:\n"
                full_prompt += "   '## Resumen' (Breve frase de lo hecho)\n"
                full_prompt += "   '## Próximos Pasos' (Qué debe hacer el usuario ahora)\n"
                full_prompt += "   [[NEXT_ACTION:{\"type\":\"...\",\"content\":\"...\"}]]\n"
            full_prompt += "ASSISTANT:"
        else:
            full_prompt = request.prompt
            if not is_analysis_mode and _is_code_task(request.prompt):
                full_prompt += "\n\n**IMPORTANTE**: 1. Genera el CÓDIGO. 2. Añade '## Resumen y Próximos Pasos'. 3. Añade `[[NEXT_ACTION:{...}]]`."

        # Solo incluir code_context en modo programación (no contaminar análisis de docs/imágenes)
        if request.code_context and not is_analysis_mode:
            full_prompt += f"\n\nContext:\n```\n{request.code_context}\n```"

        # ─── Inyectar documentos adjuntos como contexto ───────────────────────
        if request.documents:
            doc_context = "\n\n--- DOCUMENTOS ADJUNTOS ---\n"
            for doc in request.documents:
                name = doc.get("name", "documento")
                text = doc.get("text", "")
                if text.strip():
                    doc_context += f"\n📄 [{name}]:\n{text}\n"
            doc_context += "\n--- FIN DOCUMENTOS ---\n"
            full_prompt += doc_context
            logger.info(f"[CodeLab] Documentos inyectados: {[d.get('name') for d in request.documents]}")

        # ─── Intentar con modelos LAN en orden de prioridad ───────────────────
        # Máximo 2 rondas por modelo LAN (sin esperas largas para no bloquear la UI)
        MAX_ATTEMPTS_PER_MODEL = 2
        response = None
        used_model = None

        for model_config in lan_models:
            model_name = model_config.get('name', 'Unknown')
            model_id = model_config.get('id', '')

            for attempt in range(1, MAX_ATTEMPTS_PER_MODEL + 1):
                logger.info(
                    f"[CodeLab] Intentando {model_name} "
                    f"(intento {attempt}/{MAX_ATTEMPTS_PER_MODEL})"
                )
                response = await _try_lan_model(
                    model_config=model_config,
                    prompt=full_prompt,
                    system_prompt=request.system_prompt,
                    attempt=attempt,
                    images=request.images or None
                )
                if response:
                    used_model = model_id
                    break

            if response:
                break

            logger.warning(f"[CodeLab] {model_name} no disponible, probando siguiente modelo LAN...")

        if not response:
            logger.error("[CodeLab] Todos los modelos LAN fallaron.")
            logger.error(f"[CodeLab] Rutas probadas: Tailscale (100.73.44.46), LAN IP (10.13.79.31), mDNS (jddcia.local)")
            raise HTTPException(
                status_code=503,
                detail=(
                    "La IA local (Qwen3 30B) no está disponible. "
                    "Verifica que el servidor esté encendido y accesible en 10.13.79.31 o jddcia.local. "
                    "Este sistema funciona SOLO con la IA local de red LAN (sin internet)."
                )
            )

        logger.info(f"[CodeLab] ✅ Respuesta generada con modelo: {used_model}")

        # ── Post-procesado ────────────────────────────────────────────────────────

        def _extract_first_json_object(s: str) -> Optional[str]:
            """Devuelve el primer objeto JSON balanceado {…} de la cadena, o None."""
            depth = 0
            in_str = False
            escaped = False
            for i, ch in enumerate(s):
                if escaped:
                    escaped = False
                    continue
                if ch == '\\' and in_str:
                    escaped = True
                    continue
                if ch == '"':
                    in_str = not in_str
                    continue
                if in_str:
                    continue
                if ch == '{':
                    depth += 1
                elif ch == '}':
                    depth -= 1
                    if depth == 0:
                        return s[:i + 1]
            return None

        def _fix_backslashes(s: str) -> str:
            """Escapa backslashes de Windows en cadenas JSON (\\Users → \\\\Users)."""
            return re.sub(r'\\(?!["\\\/bfnrtu])', r'\\\\', s)

        def _fix_next_action(text: str) -> str:
            """Garantiza que [[NEXT_ACTION:...]] tenga JSON válido.

            Capas de reparación (en orden):
              1. Parsear como está.
              2. Extraer primer objeto JSON balanceado (descarta trailing text).
              3. Escapar backslashes Windows y reintentar.
              4. Extraer rutas del texto visible y reconstruir (fallback).
            """
            action_match = re.search(r'\[\[\s*NEXT_ACTION\s*:\s*([\s\S]*?)\]\]', text, re.IGNORECASE)
            if not action_match:
                # Sin bloque → intentar añadir uno más abajo en el flujo
                pass

            raw_inner = action_match.group(1).strip() if action_match else ""

            # Capa 1: parsear tal cual
            fixed_action_json: Optional[str] = None
            try:
                json.loads(raw_inner)
                return text  # ya correcto
            except (ValueError, KeyError):
                pass

            # Capa 2: extraer primer objeto JSON balanceado
            try:
                candidate = _extract_first_json_object(raw_inner)
                if candidate:
                    json.loads(candidate)
                    fixed_action_json = candidate
            except (ValueError, KeyError):
                pass

            # Capa 3: escapar backslashes y reintentar
            if not fixed_action_json:
                try:
                    fixed_str = _fix_backslashes(raw_inner)
                    candidate = _extract_first_json_object(fixed_str) or fixed_str
                    json.loads(candidate)
                    fixed_action_json = candidate
                except (ValueError, KeyError):
                    pass

            if fixed_action_json:
                proper_action = f"[[NEXT_ACTION:{fixed_action_json}]]"
                if action_match:
                    text = text[:action_match.start()] + proper_action + text[action_match.end():]
                    logger.info("[CodeLab] [[NEXT_ACTION]] JSON reparado (capa 2/3).")
                return text

            # Capa 4: extraer rutas Windows del texto visible y reconstruir
            # Extraer rutas Windows absolutas del texto visible (excluir el interior del NEXT_ACTION)
            text_to_scan = text[:action_match.start()] if action_match else text
            raw_paths = re.findall(r'[A-Za-z]:[/\\][^\n\r`"\'<>|*?\[\]]+', text_to_scan)
            paths = []
            seen: set = set()
            for p in raw_paths:
                p = p.strip().rstrip('/\\`"\' ,;')
                p = p.replace('/', '\\')
                if p and p not in seen and len(p) > 5:
                    seen.add(p)
                    paths.append(p)

            if not paths:
                return text  # sin rutas → no podemos generar nada

            # Comando compacto: New-Item acepta lista de paths separados por coma
            paths_str = "','".join(paths[:40])  # máx 40 para no exceder tokens de ejecución
            cmd = f"New-Item -ItemType Directory -Force -Path '{paths_str}'"
            action_json = json.dumps({"type": "run_command", "content": cmd, "label": "Crear carpetas"}, ensure_ascii=False)
            proper_action = f"[[NEXT_ACTION:{action_json}]]"

            if action_match:
                # Reemplazar el [[NEXT_ACTION:...]] malformado
                text = text[:action_match.start()] + proper_action + text[action_match.end():]
                logger.info(f"[CodeLab] [[NEXT_ACTION]] malformado corregido ({len(paths)} rutas).")
            else:
                # Añadir al final
                text = text.rstrip() + "\n\n" + proper_action
                logger.info(f"[CodeLab] [[NEXT_ACTION]] ausente — generado desde {len(paths)} rutas del texto.")

            return text

        # 1. Conversacional: quitar estructura innecesaria
        if not is_analysis_mode and not _is_code_task(request.prompt):
            response = _strip_code_structure(response)

        # 2. Tarea de código: garantizar [[NEXT_ACTION]] válido
        elif _is_code_task(request.prompt) and not is_analysis_mode:
            response = _fix_next_action(response)

            # 2b. Si aún no hay [[NEXT_ACTION]] válido (respuesta sin rutas ni código),
            #     hacer llamada de recuperación al modelo.
            still_missing = not bool(re.search(r'\[\[\s*NEXT_ACTION\s*:\s*\{', response, re.IGNORECASE))
            if still_missing and used_model:
                logger.warning("[CodeLab] [[NEXT_ACTION]] aún ausente — llamada de recuperación.")
                model_config = next((m for m in lan_models if m.get('id') == used_model), None)
                if model_config:
                    recovery_prompt = (
                        f"El usuario pidió: {request.prompt}\n\n"
                        "Tu respuesta anterior no incluyó el botón de acción [[NEXT_ACTION]].\n"
                        f"Ruta base: `{request.project_path or '.'}`\n"
                        "Genera ÚNICAMENTE una línea con el botón en formato JSON estricto:\n"
                        '[[NEXT_ACTION:{"type":"run_command","content":"<comando PowerShell>","label":"Ejecutar"}]]'
                    )
                    recovery = await _try_lan_model(
                        model_config=model_config,
                        prompt=recovery_prompt,
                        system_prompt='Responde SOLO con [[NEXT_ACTION:{...}]] en formato JSON. Nada más.',
                        attempt=1,
                        images=None
                    )
                    if recovery:
                        response = response.rstrip() + "\n\n" + recovery.strip()
                        logger.info("[CodeLab] [[NEXT_ACTION]] recuperado por llamada secundaria.")

        return {"response": response, "model": used_model}

    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"[CodeLab] Error inesperado: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8002))
    logger.info(f"[CodeLab] Iniciando backend en puerto {port}")
    logger.info(f"[CodeLab] Modo: AI_LOCAL_ONLY — Solo Qwen3 LAN (sin internet)")
    logger.info(f"[CodeLab] Modelos LAN esperados: {list(LOCAL_MODEL_IDS)}")
    uvicorn.run(app, host="127.0.0.1", port=port)
